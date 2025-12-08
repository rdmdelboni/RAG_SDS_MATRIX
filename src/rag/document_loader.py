"""Multi-format document loader for RAG knowledge base."""

from __future__ import annotations

from pathlib import Path

from langchain_core.documents import Document

from ..config.constants import SUPPORTED_FORMATS
from ..utils.logger import get_logger

logger = get_logger(__name__)


class DocumentLoader:
    """Load documents from various formats."""

    def load_file(self, file_path: Path) -> list[Document]:
        """Load a single document file.

        Args:
            file_path: Path to the file

        Returns:
            List of Document objects
        """
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        suffix = file_path.suffix.lower()
        if suffix not in SUPPORTED_FORMATS:
            raise ValueError(f"Unsupported format: {suffix}")

        logger.info("Loading document: %s", file_path.name)

        if suffix == ".pdf":
            return self._load_pdf(file_path)
        elif suffix in (".txt",):
            return self._load_text(file_path)
        elif suffix in (".md", ".markdown"):
            return self._load_markdown(file_path)
        elif suffix in (".html", ".htm"):
            return self._load_html(file_path)
        elif suffix == ".csv":
            return self._load_csv(file_path)
        elif suffix in (".docx", ".doc"):
            return self._load_docx(file_path)
        elif suffix in (".xlsx", ".xls"):
            return self._load_excel(file_path)
        else:
            raise ValueError(f"No loader for: {suffix}")

    def _load_pdf(self, file_path: Path) -> list[Document]:
        """Load PDF file."""
        try:
            import pdfplumber
        except ImportError:
            raise ImportError("pdfplumber required for PDF support")

        documents = []
        try:
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    text = page.extract_text() or ""
                    if text.strip():
                        documents.append(
                            Document(
                                page_content=text,
                                metadata={
                                    "source": str(file_path),
                                    "title": file_path.stem,
                                    "page": page_num,
                                    "type": "pdf",
                                },
                            )
                        )
        except Exception as e:
            logger.error("PDF loading failed: %s", e)
            raise

        if documents:
            logger.info("Loaded %d pages from PDF", len(documents))
            return documents

        # Fallback: use docTR when pdfplumber returns no pages/text
        doctr_docs = self._load_pdf_doctr(file_path)
        if doctr_docs:
            logger.info(
                "Loaded %d pages from PDF via docTR fallback", len(doctr_docs)
            )
            return doctr_docs

        logger.warning("No content extracted from PDF: %s", file_path)
        return []

    def _load_pdf_doctr(self, file_path: Path) -> list[Document]:
        """Load PDF using docTR directly (layout-aware OCR)."""
        try:
            from doctr.io import DocumentFile
            from doctr.models import ocr_predictor
        except ImportError:
            logger.debug("docTR not installed; skipping docTR PDF load")
            return []

        try:
            model = getattr(self, "_doctr_model", None)
            if model is None:
                logger.info("Initializing docTR model for PDF fallback")
                model = ocr_predictor(pretrained=True)
                self._doctr_model = model

            doc = DocumentFile.from_pdf(file_path)
            result = model(doc)

            docs: list[Document] = []
            for page_idx, page in enumerate(result.pages, 1):
                lines = []
                for block in page.blocks:
                    for line in block.lines:
                        line_text = " ".join(word.value for word in line.words)
                        if line_text.strip():
                            lines.append(line_text)
                page_text = "\n".join(lines)
                if page_text.strip():
                    docs.append(
                        Document(
                            page_content=page_text,
                            metadata={
                                "source": str(file_path),
                                "title": file_path.stem,
                                "page": page_idx,
                                "type": "pdf",
                                "ocr": "doctr",
                            },
                        )
                    )

            return docs
        except Exception as e:  # pragma: no cover
            logger.debug("docTR PDF fallback failed: %s", e)
            return []

    def _load_text(self, file_path: Path) -> list[Document]:
        """Load plain text file."""
        try:
            text = file_path.read_text(encoding="utf-8")
            return [
                Document(
                    page_content=text,
                    metadata={
                        "source": str(file_path),
                        "title": file_path.stem,
                        "type": "text",
                    },
                )
            ]
        except Exception as e:
            logger.error("Text loading failed: %s", e)
            raise

    def _load_markdown(self, file_path: Path) -> list[Document]:
        """Load Markdown file."""
        try:
            text = file_path.read_text(encoding="utf-8")
            return [
                Document(
                    page_content=text,
                    metadata={
                        "source": str(file_path),
                        "title": file_path.stem,
                        "type": "markdown",
                    },
                )
            ]
        except Exception as e:
            logger.error("Markdown loading failed: %s", e)
            raise

    def _load_html(self, file_path: Path) -> list[Document]:
        """Load HTML file."""
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            raise ImportError("beautifulsoup4 required for HTML support")

        try:
            html = file_path.read_text(encoding="utf-8")
            soup = BeautifulSoup(html, "html.parser")

            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()

            text = soup.get_text(separator="\n", strip=True)

            return [
                Document(
                    page_content=text,
                    metadata={
                        "source": str(file_path),
                        "title": file_path.stem,
                        "type": "html",
                    },
                )
            ]
        except Exception as e:
            logger.error("HTML loading failed: %s", e)
            raise

    def _load_csv(self, file_path: Path) -> list[Document]:
        """Load CSV file."""
        try:
            import csv
        except ImportError:
            raise ImportError("csv module required")

        try:
            with file_path.open("r", encoding="utf-8") as f:
                reader = csv.DictReader(f)
                rows = list(reader)

            # Convert rows to text
            text_parts = []
            for row in rows:
                row_text = " | ".join(f"{k}: {v}" for k, v in row.items() if v)
                text_parts.append(row_text)

            content = "\n".join(text_parts)

            return [
                Document(
                    page_content=content,
                    metadata={
                        "source": str(file_path),
                        "title": file_path.stem,
                        "type": "csv",
                        "rows": len(rows),
                    },
                )
            ]
        except Exception as e:
            logger.error("CSV loading failed: %s", e)
            raise

    def _load_docx(self, file_path: Path) -> list[Document]:
        """Load Word document."""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError("python-docx required for DOCX support")

        try:
            doc = DocxDocument(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            content = "\n".join(paragraphs)

            return [
                Document(
                    page_content=content,
                    metadata={
                        "source": str(file_path),
                        "title": file_path.stem,
                        "type": "docx",
                    },
                )
            ]
        except Exception as e:
            logger.error("DOCX loading failed: %s", e)
            raise

    def _load_excel(self, file_path: Path) -> list[Document]:
        """Load Excel file (XLSX/XLS)."""
        try:
            import openpyxl
        except ImportError:
            raise ImportError("openpyxl required for Excel support")

        try:
            documents = []
            wb = openpyxl.load_workbook(file_path, data_only=True)

            for sheet_idx, sheet_name in enumerate(wb.sheetnames, 1):
                ws = wb[sheet_name]
                rows_data = []

                # Extract all rows with data
                for row in ws.iter_rows(values_only=True):
                    if any(cell is not None for cell in row):
                        # Convert row to text
                        row_text = " | ".join(
                            str(cell).strip() for cell in row if cell is not None
                        )
                        rows_data.append(row_text)

                if rows_data:
                    content = "\n".join(rows_data)
                    documents.append(
                        Document(
                            page_content=content,
                            metadata={
                                "source": str(file_path),
                                "title": file_path.stem,
                                "type": "excel",
                                "sheet": sheet_name,
                                "sheet_number": sheet_idx,
                                "rows": len(rows_data),
                            },
                        )
                    )

            logger.info("Loaded %d sheets from Excel", len(documents))
            return documents
        except Exception as e:
            logger.error("Excel loading failed: %s", e)
            raise

    def load_directory(self, directory: Path) -> list[Document]:
        """Load all supported documents from a directory.

        Args:
            directory: Directory path

        Returns:
            List of loaded documents
        """
        directory = Path(directory)

        if not directory.is_dir():
            raise ValueError(f"Not a directory: {directory}")

        documents = []
        for file_path in sorted(directory.rglob("*")):
            if not file_path.is_file():
                continue

            if file_path.suffix.lower() not in SUPPORTED_FORMATS:
                continue

            try:
                docs = self.load_file(file_path)
                documents.extend(docs)
            except Exception as e:
                logger.warning("Failed to load %s: %s", file_path, e)
                continue

        logger.info("Loaded %d documents from directory", len(documents))
        return documents
